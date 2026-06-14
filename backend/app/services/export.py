import json
from io import BytesIO

from sqlalchemy.orm import Session
from docx import Document as DocxDocument

from app.models import Brief, Paper


def _papers_for_brief(db: Session, brief: Brief) -> list[Paper]:
    ids = {entry["paper_id"] for entry in brief.citation_map.values()}
    return db.query(Paper).filter(Paper.id.in_(ids)).all() if ids else []


def export_markdown(brief: Brief) -> str:
    lines = ["# Research Brief", "", "## Executive Summary", brief.executive_summary, ""]
    for section in brief.sections:
        lines.extend([f"## {section['theme']} ({section['consensus_status']})", ""])
        for claim in section.get("claims", []):
            refs = " ".join(f"[{ref}]" for ref in claim.get("citations", []))
            lines.append(f"- {claim['text']} {refs}")
        lines.append("")
    lines.append("## Source Passages")
    for key, cite in brief.citation_map.items():
        if str(key).startswith("__"):
            continue
        lines.append(f"- [{key}] {cite['paper_title']}, {cite['section']}, p. {cite['page']}: {cite['support_excerpt']}")
    return "\n".join(lines)


def export_csl_json(db: Session, brief: Brief) -> str:
    items = []
    for paper in _papers_for_brief(db, brief):
        items.append(
            {
                "id": str(paper.id),
                "type": "article-journal",
                "title": paper.title,
                "author": [{"literal": author.get("name", "")} for author in paper.authors],
                "issued": {"date-parts": [[paper.year]]} if paper.year else None,
                "DOI": paper.doi,
                "container-title": paper.venue,
            }
        )
    return json.dumps(items, indent=2)


def export_bibtex(db: Session, brief: Brief) -> str:
    entries = []
    for paper in _papers_for_brief(db, brief):
        key = f"paper{paper.id}"
        authors = " and ".join(author.get("name", "") for author in paper.authors)
        entries.append(
            "@article{"
            f"{key},\n"
            f"  title = {{{paper.title}}},\n"
            f"  author = {{{authors}}},\n"
            f"  year = {{{paper.year or ''}}},\n"
            f"  journal = {{{paper.venue or ''}}},\n"
            f"  doi = {{{paper.doi or ''}}}\n"
            "}"
        )
    return "\n\n".join(entries)


def export_ris(db: Session, brief: Brief) -> str:
    rows: list[str] = []
    for paper in _papers_for_brief(db, brief):
        rows.extend(["TY  - JOUR", f"TI  - {paper.title}"])
        for author in paper.authors:
            rows.append(f"AU  - {author.get('name', '')}")
        if paper.year:
            rows.append(f"PY  - {paper.year}")
        if paper.doi:
            rows.append(f"DO  - {paper.doi}")
        rows.append("ER  - ")
    return "\n".join(rows)


def export_docx(brief: Brief) -> BytesIO:
    document = DocxDocument()
    document.add_heading("Research Brief", level=1)
    document.add_heading("Executive Summary", level=2)
    document.add_paragraph(brief.executive_summary)
    for section in brief.sections:
        document.add_heading(f"{section['theme']} ({section['consensus_status']})", level=2)
        for claim in section.get("claims", []):
            refs = " ".join(f"[{ref}]" for ref in claim.get("citations", []))
            document.add_paragraph(f"{claim['text']} {refs}", style="List Bullet")
    document.add_heading("Source Passages", level=2)
    for key, cite in brief.citation_map.items():
        if str(key).startswith("__"):
            continue
        document.add_paragraph(
            f"[{key}] {cite['paper_title']}, {cite['section']}, p. {cite['page']}: {cite['support_excerpt']}",
            style="List Bullet",
        )
    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output
